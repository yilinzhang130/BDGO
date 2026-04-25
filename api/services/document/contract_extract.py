"""Plain-text extraction for legal contract files.

PDF goes through ``services.external.pubmed.extract_pdf_text`` (pypdf).
DOCX uses python-docx directly. Other extensions return empty string —
callers must check and surface a user-facing error.
"""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_contract_text(filepath: Path) -> str:
    """Extract plain text from ``filepath``. Returns empty string on failure.

    Supported: ``.pdf`` and ``.docx``. Tables in DOCX are flattened to
    pipe-separated rows so contract definition tables (``Article X | meaning``)
    survive into the LLM context.
    """
    if not filepath.exists():
        logger.warning("Contract file does not exist: %s", filepath)
        return ""
    suffix = filepath.suffix.lower()
    if suffix == ".pdf":
        from services.external.pubmed import extract_pdf_text

        return extract_pdf_text(filepath)
    if suffix == ".docx":
        return _extract_docx(filepath)
    logger.warning("Unsupported contract file type: %s", suffix)
    return ""


def _extract_docx(filepath: Path) -> str:
    try:
        from docx import Document

        doc = Document(str(filepath))
        parts: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        logger.warning("DOCX extraction failed for %s: %s", filepath, e)
        return ""
