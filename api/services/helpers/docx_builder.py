"""
Word document builder with investment-banking styling (Navy / Gold / Teal).

Designed for LLM-generated markdown reports — not a generic markdown parser.
Supports: H1/H2/H3, paragraphs with **bold**, bullet lists, pipe tables, blockquotes.

Ported & simplified from /Users/peter/.openclaw/workspace/buyside_report/md_to_docx.py
"""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

# ─────────────────────────────────────────────────────────────
# Style constants
# ─────────────────────────────────────────────────────────────

NAVY = RGBColor(0x1A, 0x3A, 0x5C)
GOLD = RGBColor(0xD4, 0x91, 0x0A)
TEAL = RGBColor(0x0E, 0x7C, 0x7B)
BODY = RGBColor(0x1C, 0x1C, 0x1C)
GREY = RGBColor(0x88, 0x88, 0x88)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

NAVY_HEX = "1A3A5C"
GOLD_HEX = "D4910A"
ALT_ROW_HEX = "F4F7FA"
CALLOUT_HEX = "FFF8E7"

FONT_EN = "Arial"
FONT_CN = "\u5fae\u8f6f\u96c5\u9ed1"  # 微软雅黑


# ─────────────────────────────────────────────────────────────
# Low-level helpers (port from md_to_docx.py)
# ─────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex: str) -> None:
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(
    run,
    size_pt: float = 10.5,
    color: RGBColor = BODY,
    bold: bool = False,
    italic: bool = False,
) -> None:
    """Configure a run with English + Chinese fallback fonts."""
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = FONT_EN
    # Set East Asian font so Chinese chars render 微软雅黑
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_CN)


def add_formatted_run(paragraph, text: str, size_pt: float = 10.5,
                      color: RGBColor = BODY, bold: bool = False) -> None:
    """Add text to a paragraph, parsing inline **bold** markdown."""
    if not text:
        return
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size_pt, color, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size_pt, color, bold=bold)


# ─────────────────────────────────────────────────────────────
# High-level builder API
# ─────────────────────────────────────────────────────────────

def new_report_document() -> Document:
    """Create a styled Document with A4 margins + default body font."""
    doc = Document()
    for section in doc.sections:
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    # Set default style for Normal paragraph so Chinese renders correctly
    style = doc.styles["Normal"]
    style.font.name = FONT_EN
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_CN)
    return doc


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    """Large cover-style title block (centered, navy + gold)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, size_pt=22, color=NAVY, bold=True)

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        set_run_font(run, size_pt=13, color=GOLD, bold=True)

    # Gold separator line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2500" * 40)
    set_run_font(run, size_pt=10, color=GOLD)
    doc.add_paragraph()


def add_h1(doc: Document, text: str) -> None:
    """Level-1 heading: navy 16pt with gold underline."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size_pt=16, color=NAVY, bold=True)

    # Gold bottom border on paragraph
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="8" w:space="2" w:color="{GOLD_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_h2(doc: Document, text: str) -> None:
    """Level-2 heading: navy 13pt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size_pt=13, color=NAVY, bold=True)


def add_h3(doc: Document, text: str) -> None:
    """Level-3 heading: teal 11.5pt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size_pt=11.5, color=TEAL, bold=True)


def add_body_paragraph(doc: Document, text: str) -> None:
    """Plain paragraph with **bold** support."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    add_formatted_run(p, text, size_pt=10.5, color=BODY)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_formatted_run(p, text, size_pt=10.5, color=BODY)


def add_blockquote(doc: Document, text: str) -> None:
    """Highlighted callout-style block (teal left border, body text)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(8)
    add_formatted_run(p, text, size_pt=10, color=GREY, bold=False)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="8" w:color="0E7C7B"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    """Render a markdown table. First row is the header (navy bg, white text).
    Alternating body rows get a light gray background.
    """
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(rows):
        tr = table.add_row()
        padded = list(row_data) + [""] * (num_cols - len(row_data))
        for col_idx, cell_text in enumerate(padded):
            cell = tr.cells[col_idx]
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if row_idx == 0:
                set_cell_shading(cell, NAVY_HEX)
                add_formatted_run(p, cell_text, size_pt=9.5, color=WHITE, bold=True)
            else:
                if row_idx % 2 == 0:
                    set_cell_shading(cell, ALT_ROW_HEX)
                add_formatted_run(p, cell_text, size_pt=9.5, color=BODY)

    # Full-width
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    existing_w = tblPr.find(qn("w:tblW"))
    if existing_w is None:
        tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>'))
    else:
        existing_w.set(qn("w:type"), "pct")
        existing_w.set(qn("w:w"), "5000")

    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────
# Markdown → docx renderer
# ─────────────────────────────────────────────────────────────

_TABLE_ROW = re.compile(r"^\|(.+)\|\s*$")
_TABLE_SEP = re.compile(r"^\|?[\s\-:|]+\|?\s*$")
_BULLET = re.compile(r"^[-*]\s+(.*)")


def markdown_to_docx(markdown: str, doc: Document) -> None:
    """Convert LLM-generated markdown into styled paragraphs/headings/tables.

    Scope: H1-H3, paragraphs (with inline **bold**), bullet lists, pipe tables,
    and > blockquotes. Does not handle arbitrary CommonMark (no images, no
    links rendered as hyperlinks, no code fences).
    """
    lines = markdown.splitlines()
    i = 0
    n = len(lines)

    while i < n:
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        # Skip blank lines (but respect them as paragraph separators — handled naturally)
        if not stripped:
            i += 1
            continue

        # Skip horizontal rules
        if re.match(r"^-{3,}$|^={3,}$|^\*{3,}$", stripped):
            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            add_h1(doc, stripped[2:].strip())
            i += 1
            continue
        if stripped.startswith("## "):
            add_h2(doc, stripped[3:].strip())
            i += 1
            continue
        if stripped.startswith("### "):
            add_h3(doc, stripped[4:].strip())
            i += 1
            continue
        if stripped.startswith("#### "):
            # treat as h3 (downgrade, we only have 3 levels of styling)
            add_h3(doc, stripped[5:].strip())
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            add_blockquote(doc, " ".join(quote_lines))
            continue

        # Bullets
        if _BULLET.match(stripped):
            while i < n:
                m = _BULLET.match(lines[i].strip())
                if not m:
                    break
                add_bullet(doc, m.group(1))
                i += 1
            continue

        # Tables
        if _TABLE_ROW.match(stripped):
            table_rows: list[list[str]] = []
            while i < n and _TABLE_ROW.match(lines[i].strip()):
                if _TABLE_SEP.match(lines[i].strip()):
                    i += 1
                    continue
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                table_rows.append(cells)
                i += 1
            if table_rows:
                add_table(doc, table_rows)
            continue

        # Plain paragraph — coalesce consecutive non-structural lines
        para_lines = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].rstrip()
            nxt_strip = nxt.strip()
            if (
                not nxt_strip
                or nxt_strip.startswith("#")
                or nxt_strip.startswith(">")
                or _BULLET.match(nxt_strip)
                or _TABLE_ROW.match(nxt_strip)
            ):
                break
            para_lines.append(nxt_strip)
            i += 1
        add_body_paragraph(doc, " ".join(para_lines))


def document_to_bytes(doc: Document) -> bytes:
    """Serialise a Document to bytes for ctx.save_file(format='docx')."""
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
