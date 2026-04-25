"""Tests for contract text extraction."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from services.document.contract_extract import extract_contract_text


def test_missing_file_returns_empty(tmp_path: Path):
    assert extract_contract_text(tmp_path / "does_not_exist.pdf") == ""


def test_unsupported_extension_returns_empty(tmp_path: Path):
    f = tmp_path / "contract.txt"
    f.write_text("Some text")
    assert extract_contract_text(f) == ""


def test_docx_paragraphs_extracted(tmp_path: Path):
    doc = Document()
    doc.add_paragraph("Section 1: Confidentiality.")
    doc.add_paragraph("The Receiving Party shall not disclose...")
    path = tmp_path / "contract.docx"
    doc.save(str(path))

    text = extract_contract_text(path)
    assert "Section 1: Confidentiality." in text
    assert "The Receiving Party" in text


def test_docx_table_flattened_to_pipe_rows(tmp_path: Path):
    doc = Document()
    doc.add_paragraph("Definitions table follows.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Term"
    table.cell(0, 1).text = "Meaning"
    table.cell(1, 0).text = "Material"
    table.cell(1, 1).text = "The substance described in Appendix A"
    path = tmp_path / "with_table.docx"
    doc.save(str(path))

    text = extract_contract_text(path)
    assert "Term | Meaning" in text
    assert "Material | The substance described in Appendix A" in text
